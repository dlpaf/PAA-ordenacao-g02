"""
gerar_tabelas_artigo.py
=======================
Gera tabelas comparativas entre os algoritmos clássicos e o AOH (Threshold=16)
para inclusão no artigo/relatório.

Como usar:
    python gerar_tabelas_artigo.py

Saída:
    tabelas_artigo.docx  — tabelas prontas para copiar no Word
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Dados da Parte 1 + AOH (Threshold 16) ────────────────────────────────────
TAMANHOS = [100, 1000, 5000, 30000, 50000, 100000, 150000, 200000]

TEMPOS = {
    "Bubble Sort": {
        "Crescente":   [2.66e-4, 2.60e-2, 0.6026, 21.430, 57.731, 250.52, 573.56, 1026.61],
        "Decrescente": [3.12e-4, 4.54e-2, 1.1115, 41.195, 111.33, 459.29, 1108.1, 2020.09],
        "Aleatório":   [2.72e-4, 4.02e-2, 0.9005, 33.292, 93.461, 402.92, 916.57, 1716.11],
    },
    "Insertion Sort": {
        "Crescente":   [7.87e-6, 6.77e-5, 3.77e-4, 2.28e-3, 3.75e-3, 7.40e-3, 1.14e-2, 1.60e-2],
        "Decrescente": [2.52e-4, 3.33e-2, 0.8904, 32.280, 92.724, 384.27, 867.29, 1639.29],
        "Aleatório":   [1.46e-4, 1.94e-2, 0.4532, 16.337, 45.713, 185.02, 437.93, 801.90],
    },
    "Merge Sort": {
        "Crescente":   [1.10e-4, 1.54e-3, 9.98e-3, 7.09e-2, 0.1031, 0.2221, 0.3423, 0.4714],
        "Decrescente": [1.08e-4, 1.47e-3, 1.03e-2, 5.90e-2, 0.1030, 0.2182, 0.3416, 0.4777],
        "Aleatório":   [1.08e-4, 1.63e-3, 1.05e-2, 6.22e-2, 0.1089, 0.2335, 0.3705, 0.5191],
    },
    "Heap Sort": {
        "Crescente":   [1.27e-4, 2.24e-3, 1.50e-2, 0.1093, 0.1666, 0.3541, 0.5574, 0.7849],
        "Decrescente": [1.02e-4, 1.73e-3, 1.35e-2, 8.70e-2, 0.1531, 0.3310, 0.5179, 0.7138],
        "Aleatório":   [1.15e-4, 1.94e-3, 1.41e-2, 9.37e-2, 0.1722, 0.3628, 0.5678, 0.8057],
    },
    "Quick Sort": {
        "Crescente":   [7.26e-5, 1.06e-3, 5.77e-3, 5.26e-2, 7.05e-2, 0.1461, 0.2294, 0.3111],
        "Decrescente": [6.65e-5, 1.27e-3, 6.82e-3, 3.91e-2, 7.08e-2, 0.1523, 0.2353, 0.3087],
        "Aleatório":   [6.84e-5, 1.05e-3, 7.08e-3, 4.36e-2, 7.55e-2, 0.1629, 0.2605, 0.3534],
    },
    "AOH (T=16)": {
        "Crescente":   [3.57e-5, 6.74e-4, 4.91e-3, 3.19e-2, 5.60e-2, 0.1241, 0.1848, 0.2385],
        "Decrescente": [4.17e-5, 8.75e-4, 7.31e-3, 3.82e-2, 6.87e-2, 0.1393, 0.2183, 0.3328],
        "Aleatório":   [4.38e-5, 7.35e-4, 6.55e-3, 4.60e-2, 7.48e-2, 0.1556, 0.2501, 0.3574],
    },
}

COMPARACOES = {
    "Bubble Sort": {
        "Crescente":   [4950, 499500, 12497500, 449985000, 1249975000, 4999950000, 11249925000, 19999900000],
        "Decrescente": [4950, 499500, 12497500, 449985000, 1249975000, 4999950000, 11249925000, 19999900000],
        "Aleatório":   [4950, 499500, 12497500, 449985000, 1249975000, 4999950000, 11249925000, 19999900000],
    },
    "Insertion Sort": {
        "Crescente":   [99, 999, 4999, 29999, 49999, 99999, 149999, 199999],
        "Decrescente": [4950, 499500, 12497500, 449985000, 1249975000, 4999950000, 11249925000, 19999900000],
        "Aleatório":   [2738, 252033, 6193391, 224278917, 626181777, 2509455164, 5613832908, 9973802293],
    },
    "Merge Sort": {
        "Crescente":   [672, 9976, 61808, 447232, 784464, 1668928, 2587856, 3537856],
        "Decrescente": [672, 9976, 61808, 447232, 784464, 1668928, 2587856, 3537856],
        "Aleatório":   [672, 9976, 61808, 447232, 784464, 1668928, 2587856, 3537856],
    },
    "Heap Sort": {
        "Crescente":   [1081, 17583, 112126, 826347, 1455438, 3112517, 4838906, 6626912],
        "Decrescente": [944,  15965, 103227, 775687, 1366047, 2926640, 4571170, 6256303],
        "Aleatório":   [1034, 16895, 107681, 800852, 1409969, 3019243, 4700915, 6440303],
    },
    "Quick Sort": {
        "Crescente":   [660, 11062, 66170, 535409, 937344, 2007740, 3162895, 4193823],
        "Decrescente": [641, 10888, 70078, 527901, 969865, 2063974, 3112641, 4230386],
        "Aleatório":   [591, 10930, 71513, 555815, 950219, 2011064, 3200704, 4231088],
    },
    "AOH (T=16)": {
        "Crescente":   [486, 9767, 61609, 487102, 867695, 1939823, 2953479, 3844858],
        "Decrescente": [728, 12793, 74252, 573309, 1005193, 2115875, 3289519, 4652837],
        "Aleatório":   [762, 11360, 73737, 544239, 961391, 2121571, 3265709, 4515620],
    },
}

CENARIOS = ["Crescente", "Decrescente", "Aleatório"]
ALGORITMOS = list(TEMPOS.keys())

# ─── Cores ────────────────────────────────────────────────────────────────────
COR_CABECALHO  = "1F4E79"   # azul escuro
COR_AOH        = "C6EFCE"   # verde claro  — linha do AOH
COR_LINHA_PAR  = "EBF3FB"   # azul bem claro
COR_LINHA_IMPAR= "FFFFFF"   # branco
COR_MELHOR     = "FFEB9C"   # amarelo — melhor valor da coluna


def fmt_tempo(v):
    """Formata tempo: notação científica para valores muito pequenos, fixo para grandes."""
    if v >= 100:
        return f"{v:.1f}"
    if v >= 1:
        return f"{v:.4f}"
    if v >= 0.001:
        return f"{v:.5f}"
    return f"{v:.2e}"


def fmt_comp(v):
    """Formata comparações com separador de milhar."""
    return f"{int(v):,}".replace(",", ".")


def set_cell_bg(cell, hex_color):
    """Define cor de fundo de uma célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="BBBBBB"):
    """Define bordas finas em uma célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def escrever_celula(cell, texto, negrito=False, cor_texto="000000",
                    tamanho=9, centralizar=True, bg=None):
    """Escreve texto formatado em uma célula."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        set_cell_bg(cell, bg)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centralizar else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(texto))
    run.bold = negrito
    run.font.size = Pt(tamanho)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(cor_texto)


def adicionar_tabela_tempo(doc, cenario):
    """Cria tabela de tempo médio para um cenário."""
    doc.add_paragraph()
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = titulo.add_run(f"Tabela — Tempo Médio de Execução (s) — Vetor {cenario}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string("1F4E79")

    n_cols = 1 + len(TAMANHOS)
    tabela = doc.add_table(rows=1 + len(ALGORITMOS), cols=n_cols)
    tabela.style = "Table Grid"

    # Larguras
    tabela.columns[0].width = Cm(3.2)
    for i in range(1, n_cols):
        tabela.columns[i].width = Cm(1.8)

    # Cabeçalho
    cab = tabela.rows[0].cells
    escrever_celula(cab[0], "Algoritmo", negrito=True, cor_texto="FFFFFF", bg=COR_CABECALHO)
    for i, n in enumerate(TAMANHOS):
        escrever_celula(cab[i+1], f"{n:,}".replace(",","."),
                        negrito=True, cor_texto="FFFFFF", bg=COR_CABECALHO)

    # Encontrar melhor (menor) por coluna
    melhores = []
    for i in range(len(TAMANHOS)):
        vals = [TEMPOS[alg][cenario][i] for alg in ALGORITMOS]
        melhores.append(min(vals))

    # Linhas de dados
    for ai, alg in enumerate(ALGORITMOS):
        row = tabela.rows[ai + 1].cells
        is_aoh = alg == "AOH (T=16)"
        bg_linha = COR_AOH if is_aoh else (COR_LINHA_PAR if ai % 2 == 0 else COR_LINHA_IMPAR)

        escrever_celula(row[0], alg, negrito=is_aoh,
                        cor_texto="000000", bg=bg_linha, centralizar=False)
        for i, n in enumerate(TAMANHOS):
            v = TEMPOS[alg][cenario][i]
            # Destaca melhor valor da coluna em amarelo
            bg = COR_MELHOR if v == melhores[i] else bg_linha
            escrever_celula(row[i+1], fmt_tempo(v), negrito=(v == melhores[i]), bg=bg)

    # Legenda
    leg = doc.add_paragraph()
    leg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = leg.add_run("Verde = AOH (T=16).  Amarelo = melhor tempo da coluna.  Valores em segundos (s).")
    rl.font.size = Pt(8)
    rl.font.name = "Arial"
    rl.font.color.rgb = RGBColor.from_string("666666")
    rl.italic = True


def adicionar_tabela_comparacoes(doc, cenario):
    """Cria tabela de comparações médias para um cenário."""
    doc.add_paragraph()
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = titulo.add_run(f"Tabela — Comparações Médias — Vetor {cenario}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string("1F4E79")

    n_cols = 1 + len(TAMANHOS)
    tabela = doc.add_table(rows=1 + len(ALGORITMOS), cols=n_cols)
    tabela.style = "Table Grid"

    tabela.columns[0].width = Cm(3.2)
    for i in range(1, n_cols):
        tabela.columns[i].width = Cm(1.8)

    cab = tabela.rows[0].cells
    escrever_celula(cab[0], "Algoritmo", negrito=True, cor_texto="FFFFFF", bg=COR_CABECALHO)
    for i, n in enumerate(TAMANHOS):
        escrever_celula(cab[i+1], f"{n:,}".replace(",","."),
                        negrito=True, cor_texto="FFFFFF", bg=COR_CABECALHO)

    melhores = []
    for i in range(len(TAMANHOS)):
        vals = [COMPARACOES[alg][cenario][i] for alg in ALGORITMOS]
        melhores.append(min(vals))

    for ai, alg in enumerate(ALGORITMOS):
        row = tabela.rows[ai + 1].cells
        is_aoh = alg == "AOH (T=16)"
        bg_linha = COR_AOH if is_aoh else (COR_LINHA_PAR if ai % 2 == 0 else COR_LINHA_IMPAR)

        escrever_celula(row[0], alg, negrito=is_aoh,
                        cor_texto="000000", bg=bg_linha, centralizar=False)
        for i in range(len(TAMANHOS)):
            v = COMPARACOES[alg][cenario][i]
            bg = COR_MELHOR if v == melhores[i] else bg_linha
            escrever_celula(row[i+1], fmt_comp(v), negrito=(v == melhores[i]), bg=bg)

    leg = doc.add_paragraph()
    rl = leg.add_run("Verde = AOH (T=16).  Amarelo = menor número de comparações da coluna.")
    rl.font.size = Pt(8)
    rl.font.name = "Arial"
    rl.font.color.rgb = RGBColor.from_string("666666")
    rl.italic = True


def adicionar_tabela_ganho(doc):
    """Tabela de ganho percentual do AOH sobre o Quick Sort puro."""
    doc.add_paragraph()
    titulo = doc.add_paragraph()
    r = titulo.add_run("Tabela — Ganho de Tempo do AOH (T=16) sobre o Quick Sort puro (%)")
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string("1F4E79")

    tabela = doc.add_table(rows=1 + len(TAMANHOS), cols=4)
    tabela.style = "Table Grid"
    tabela.columns[0].width = Cm(3.0)
    for i in range(1, 4):
        tabela.columns[i].width = Cm(3.0)

    cab = tabela.rows[0].cells
    for i, txt in enumerate(["Tamanho (n)", "Crescente", "Decrescente", "Aleatório"]):
        escrever_celula(cab[i], txt, negrito=True, cor_texto="FFFFFF", bg=COR_CABECALHO)

    for ri, n in enumerate(TAMANHOS):
        row = tabela.rows[ri + 1].cells
        bg = COR_LINHA_PAR if ri % 2 == 0 else COR_LINHA_IMPAR
        escrever_celula(row[0], f"{n:,}".replace(",","."), negrito=True, bg=bg)
        for ci, cen in enumerate(CENARIOS):
            q = TEMPOS["Quick Sort"][cen][ri]
            a = TEMPOS["AOH (T=16)"][cen][ri]
            ganho = (q - a) / q * 100
            txt = f"{ganho:+.1f}%"
            # Verde se ganhou, vermelho se perdeu
            cor_bg = "C6EFCE" if ganho > 0 else "FFCCCC"
            escrever_celula(row[ci + 1], txt, negrito=ganho > 0, bg=cor_bg)

    leg = doc.add_paragraph()
    rl = leg.add_run("Verde = AOH mais rápido que QuickSort.  Vermelho = QuickSort mais rápido.")
    rl.font.size = Pt(8)
    rl.font.name = "Arial"
    rl.font.color.rgb = RGBColor.from_string("666666")
    rl.italic = True


# ─── Montar o documento ───────────────────────────────────────────────────────
def main():
    doc = Document()

    # Margens
    section = doc.sections[0]
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    # Título
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("Tabelas Comparativas — AOH vs Algoritmos Clássicos")
    rt.bold = True
    rt.font.size = Pt(14)
    rt.font.name = "Arial"
    rt.font.color.rgb = RGBColor.from_string("1F4E79")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Algoritmo de Ordenação Híbrido: QuickSort (pivô aleatório) + Insertion Sort | Threshold = 16")
    rs.font.size = Pt(10)
    rs.font.name = "Arial"
    rs.font.color.rgb = RGBColor.from_string("444444")
    rs.italic = True

    # Tabelas de tempo por cenário
    for cen in CENARIOS:
        adicionar_tabela_tempo(doc, cen)

    doc.add_page_break()

    # Tabelas de comparações por cenário
    for cen in CENARIOS:
        adicionar_tabela_comparacoes(doc, cen)

    doc.add_page_break()

    # Tabela de ganho percentual
    adicionar_tabela_ganho(doc)

    saida = "tabelas_artigo.docx"
    doc.save(saida)
    print(f"✓ Arquivo gerado: {saida}")
    print(f"  Tabelas incluídas:")
    print(f"    - Tempo médio (s): Crescente, Decrescente, Aleatório")
    print(f"    - Comparações médias: Crescente, Decrescente, Aleatório")
    print(f"    - Ganho % do AOH sobre o QuickSort puro")


if __name__ == "__main__":
    main()